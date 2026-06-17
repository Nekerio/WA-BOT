import io
import logging
import docx

logger = logging.getLogger(__name__)

def extract_text_from_doc(doc_bytes: bytes, filename: str) -> str:
    """
    Mengekstrak teks dari file Word (.docx atau .doc).
    Jika file adalah .docx, kita baca menggunakan python-docx.
    Jika file adalah .doc (format lama), minta pengguna menyimpannya ke .docx.
    """
    if filename.lower().endswith(".doc") and not filename.lower().endswith(".docx"):
        return "[ERROR] Format .doc (Word lama) tidak didukung langsung. Silakan simpan dokumen Anda sebagai .docx (Word Document) lalu unggah kembali."
        
    try:
        doc = docx.Document(io.BytesIO(doc_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Baca teks dari tabel jika ada
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        result_text = '\n'.join(full_text)
        if not result_text.strip():
            return "[ERROR] Dokumen Word kosong atau tidak memiliki teks yang terbaca."
            
        return result_text
    except Exception as e:
        logger.error(f"Gagal mengekstrak teks dari file Word: {e}")
        return f"[ERROR] Gagal membaca dokumen Word: {str(e)}"
